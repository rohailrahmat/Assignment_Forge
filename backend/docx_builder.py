from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re
from datetime import datetime

# Premium Color Palette
NAVY_BLUE = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT_BLUE = RGBColor(0x63, 0x82, 0xFF)
SLATE_GREY = RGBColor(0x47, 0x55, 0x69)
LIGHT_GREY = RGBColor(0xF8, 0xF9, 0xFA)
BORDER_COL = RGBColor(0xDD, 0xDD, 0xDD)

def set_cell_border(cell, **kwargs):
    """
    Set cell borders for a table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def build_docx(content: dict, output_path: str, student_name: str, course_code: str):
    doc = Document()
    
    # Set Margins (Standard Academic: 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Apply Standard Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # 1. COVER PAGE
    add_executive_cover_page(doc, content, student_name, course_code)
    doc.add_page_break()

    # 2. TABLE OF CONTENTS (Manual)
    add_table_of_contents(doc, content)
    doc.add_page_break()

    # 3. CONTENT SECTIONS
    for i, section_data in enumerate(content.get("sections", []), 1):
        add_section(doc, section_data, i)
        
    # 4. REFERENCES SECTION (Academic Standard)
    add_references_placeholder(doc, content)

    # 5. FOOTER & HEADER
    add_headers_footers(doc, student_name, course_code)

    doc.save(str(output_path))
    return str(output_path)

def add_executive_cover_page(doc, content, student_name, course_code):
    doc.add_paragraph("\n" * 2)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(content.get("title", "ACADEMIC SOLUTION").upper())
    run.font.name = 'Georgia'
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = NAVY_BLUE

    doc.add_paragraph("\n" * 1)
    
    # Subtitle/Course
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(f"Course: {content.get('course', course_code)}")
    run.font.name = 'Georgia'
    run.font.size = Pt(16)
    run.font.color.rgb = SLATE_GREY

    doc.add_paragraph("\n" * 4)

    # Metadata Table
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    meta_items = [
        ("Student Name", student_name),
        ("Course Code", content.get("course", course_code)),
        ("Instructor", content.get("instructor", "TBD")),
        ("Assignment No.", content.get("assignment_number", "1")),
        ("Submission Date", datetime.now().strftime("%B %d, %Y")),
        ("Status", "Final Draft - Executive Grade")
    ]

    for label, value in meta_items:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        
        # Style label
        p0 = row[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.runs[0]
        r0.bold = True
        r0.font.name = 'Georgia'
        r0.font.color.rgb = NAVY_BLUE
        
        # Style value
        p1 = row[1].paragraphs[0]
        r1 = p1.runs[0]
        r1.font.name = 'Times New Roman'
        r1.font.color.rgb = SLATE_GREY

def add_table_of_contents(doc, content):
    p = doc.add_paragraph("TABLE OF CONTENTS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = NAVY_BLUE
    
    doc.add_paragraph("\n")
    
    for i, section in enumerate(content.get("sections", []), 1):
        row_p = doc.add_paragraph()
        run = row_p.add_run(f"{i}. {section.get('heading', 'Section').upper()}")
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.font.color.rgb = SLATE_GREY

def add_section(doc, section_data, index):
    # Section Header
    header_p = doc.add_paragraph()
    header_p.space_before = Pt(24)
    run = header_p.add_run(f"{index}. {section_data.get('heading', '').upper()}")
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = NAVY_BLUE

    # Visual Divider (More subtle)
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(6)
    run = divider.add_run("_" * 40)
    run.font.color.rgb = BORDER_COL
    run.bold = True

    # Content
    for item in section_data.get("content", []):
        render_content_item(doc, item)

def render_content_item(doc, item):
    itype = item.get("type")
    
    if itype == "paragraph":
        text = str(item.get("text", ""))
        label = item.get("label", "")
        
        # Check for image[[...]]
        if "image[[" in text:
            add_image_placeholder(doc, text)
            return

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(12)
        
        if label:
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.name = 'Georgia'
            run.font.color.rgb = NAVY_BLUE
            
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    elif itype == "field":
        p = doc.add_paragraph()
        run = p.add_run(f"{item.get('label', '')}: ")
        run.bold = True
        run.font.color.rgb = NAVY_BLUE
        p.add_run(str(item.get('value', '')))

    elif itype in ["bullet_list", "numbered_list"]:
        if item.get("label"):
            p = doc.add_paragraph(item.get("label"))
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = NAVY_BLUE
            
        for li in item.get("items", []):
            p = doc.add_paragraph(str(li), style='List Number' if itype == "numbered_list" else 'List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)

    elif itype == "grid_table":
        headers = item.get("headers", [])
        rows = item.get("rows", [])
        if not headers: return
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h).upper()
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.name = 'Georgia'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Background color via XML
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '1A1A2E')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        # Data Rows
        for idx, r in enumerate(rows):
            row_cells = table.add_row().cells
            for i, val in enumerate(r):
                row_cells[i].text = str(val)
                # Alternate row shading
                if idx % 2 == 0:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F8F9FA')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                
                p = row_cells[i].paragraphs[0]
                run = p.runs[0] if p.runs else p.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    elif itype == "pros_cons_table":
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr = table.rows[0].cells
        hdr[0].text = "STRATEGIC ADVANTAGES (PROS)"
        hdr[1].text = "POTENTIAL LIMITATIONS (CONS)"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = NAVY_BLUE
            # Add header background
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F1F5F9')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
        pros = item.get("pros", [])
        cons = item.get("cons", [])
        for i in range(max(len(pros), len(cons))):
            row = table.add_row().cells
            row[0].text = pros[i] if i < len(pros) else ""
            row[1].text = cons[i] if i < len(cons) else ""

def add_image_placeholder(doc, text):
    import re
    match = re.search(r'image\[\[(.*?)\]\]', text)
    desc = match.group(1) if match else "Strategic Visualization"
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_border(cell, top={"sz": 6, "val": "dashed", "color": "6382FF"},
                          bottom={"sz": 6, "val": "dashed", "color": "6382FF"},
                          start={"sz": 6, "val": "dashed", "color": "6382FF"},
                          end={"sz": 6, "val": "dashed", "color": "6382FF"})
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n📷 [ FIGURE PLACEHOLDER ]\n")
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = NAVY_BLUE
    
    p2 = cell.add_paragraph(desc.upper())
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run()
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = SLATE_GREY
    
    # Add Hint
    p3 = cell.add_paragraph(f"\n💡 PRO TIP: {text.split('hint:')[-1] if 'hint:' in text else 'Capture the full dashboard with browser address bar for maximum authenticity.'}")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.runs[0]
    run3.font.name = 'Georgia'
    run3.font.size = Pt(9)
    run3.font.color.rgb = ACCENT_BLUE
    run3.italic = True

def add_references_placeholder(doc, content):
    doc.add_page_break()
    p = doc.add_paragraph("REFERENCES")
    run = p.add_run()
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = NAVY_BLUE
    
    doc.add_paragraph("\n[ 1 ] Academic standards require at least 3-5 peer-reviewed sources here.")
    doc.add_paragraph("[ 2 ] Industry reports and whitepapers from verified sources.")
    doc.add_paragraph("[ 3 ] Contextual data extracted from associated entities.")

def add_headers_footers(doc, student_name, course_code):
    for section in doc.sections:
        # Header
        header = section.header
        p = header.paragraphs[0]
        p.text = f"Academic Solution | {student_name} | {course_code}"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = SLATE_GREY
        
        # Footer
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"© {datetime.now().year} AssignmentForge Elite - Confidential Research Document"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = SLATE_GREY
